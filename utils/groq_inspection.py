"""Generate a highly detailed vehicle inspection narrative via Groq LLM and export to Word (.docx)."""

from __future__ import annotations

import os
import re
from datetime import datetime
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

try:
    from groq import Groq
except ImportError:  # pragma: no cover
    Groq = None  # type: ignore


NA = "Not available from free sources"

# Brand colours
NAVY = RGBColor(0x1E, 0x3A, 0x5F)
STEEL = RGBColor(0x2C, 0x52, 0x7A)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0x88, 0x88, 0x88)
BLACK = RGBColor(0x22, 0x22, 0x22)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HDR_BG = "1E3A5F"
ROW_ALT_BG = "F4F6F8"


def _v(val: Any, fallback: str = NA) -> str:
    if val is None:
        return fallback
    s = str(val).strip()
    return s if s else fallback


def _build_context(report: dict) -> str:
    """Flatten the VIN report into a clean context string for the LLM."""
    ident = report.get("identification") or {}
    engine = report.get("engine") or {}
    drive = report.get("drivetrain") or {}
    dims = report.get("dimensions") or {}
    mfg = report.get("manufacturing") or {}
    safety = report.get("safety") or {}
    manual = report.get("manual") or {}
    wmi = report.get("wmi") or {}
    cond = manual.get("condition") or {}

    recalls_text = []
    for r in (safety.get("recalls") or [])[:4]:
        summary = (r.get("summary") or "")[:120]
        recalls_text.append(
            f"- {r.get('campaign_number')}: {r.get('component') or 'N/A'} — {summary}"
        )
    recalls_block = "\n".join(recalls_text) if recalls_text else "No recalls for Year/Make/Model."

    notes = (_v(manual.get("notes"), "None") or "None")[:400]
    findings = (_v(manual.get("additional_findings"), "None") or "None")[:300]

    ctx = f"""VIN: {report.get('vin')}
Report: {report.get('id')} | Completeness: {report.get('data_completeness', 0)}%

IDENTIFICATION
Make/Model/Year: {_v(ident.get('make'))} {_v(ident.get('model'))} {_v(ident.get('model_year'))}
Trim/Series: {_v(ident.get('trim'))} / {_v(ident.get('series'))}
Body/Type/Class: {_v(ident.get('body_class'))} / {_v(ident.get('vehicle_type'))} / {_v(ident.get('vehicle_class'))}
Manufacturer: {_v(ident.get('manufacturer'))}

ENGINE & DRIVETRAIN
Engine: {_v(engine.get('engine_model'))} | {_v(engine.get('displacement'))} | {_v(engine.get('cylinders'))} cyl | {_v(engine.get('configuration'))}
Fuel/HP: {_v(engine.get('fuel_type'))} | {_v(engine.get('horsepower'))} hp | Mfr: {_v(engine.get('manufacturer'))}
Trans/Drive: {_v(drive.get('transmission'))} ({_v(drive.get('transmission_speeds'))} spd) | {_v(drive.get('drive_type'))}
Doors/Seats/GVWR: {_v(dims.get('doors'))} / {_v(dims.get('seats'))} / {_v(dims.get('gvwr'))}

MANUFACTURING
Plant: {_v(mfg.get('plant_city'))}, {_v(mfg.get('plant_state'))}, {_v(mfg.get('plant_country'))} ({_v(mfg.get('plant_company'))})
WMI: {_v(wmi.get('code'))} | {_v(wmi.get('manufacturer'))} | {_v(wmi.get('country'))}

RECALLS (Year/Make/Model — verify VIN at nhtsa.gov/recalls)
{recalls_block}

INSPECTOR DATA
Client: {_v(manual.get('client_name'), '—')} | Ref: {_v(manual.get('client_reference'), '—')}
Reg/Mileage/Price: {_v(manual.get('registration'), '—')} | {_v(manual.get('mileage'), '—')} | {_v(manual.get('purchase_price'), '—')}
Date/Inspector: {_v(manual.get('inspection_date'), '—')} | {_v(manual.get('inspector_name'), '—')}
Condition overall/ext/int/eng/trans/elec/tires/brakes: {cond.get('overall','n/a')} / {cond.get('exterior','n/a')} / {cond.get('interior','n/a')} / {cond.get('engine','n/a')} / {cond.get('transmission','n/a')} / {cond.get('electrical','n/a')} / {cond.get('tires','n/a')} / {cond.get('brakes','n/a')}
Notes: {notes}
Findings: {findings}
"""
    return ctx.strip()


SYSTEM_PROMPT = """You are an expert automotive inspector and technical writer.
Write a detailed, professional vehicle inspection report from the supplied VIN data only.

RULES:
- Never invent accidents, ownership, title brands, odometer, or service history.
- If data is missing, write: Not available from free public sources.
- Treat recalls as Year/Make/Model campaigns; remind to verify exact VIN at nhtsa.gov/recalls.
- Incorporate inspector notes/ratings when present; label them as inspector observations.
- Use clear professional English.
- Do NOT use code fences, markdown tables, or escaped characters like \\*.
- Use plain ## headings, - for bullets, and **bold** only for short labels.
- Do not repeat the report title or VIN at the top of the body.

Structure EXACTLY with these ## headings:
## Executive Summary
## Vehicle Identification
## Technical Specifications
## Manufacturing Origin
## Safety & Recall Analysis
## Condition Assessment (Inspector Observations)
## Recommended Inspection Checklist
## Known Common Issues for this Model
## Limitations & Disclaimer
## Closing Recommendation

Checklist: comprehensive bullets for exterior, underbody, engine bay, interior, fluids, electronics, tires/brakes, road test — tailored to class/age.
Common issues: only well-known patterns for that generation, or say research is needed.
Be thorough and factual.
"""


# Preferred order when preferred model fails (rate limit, 413, decommissioned, etc.)
# Updated Sep 2026 — older Llama 3.x / Mixtral / Gemma IDs are decommissioned on free/dev tier.
# See: https://console.groq.com/docs/deprecations
FALLBACK_MODELS = [
    "openai/gpt-oss-120b",
    "openai/gpt-oss-20b",
    "qwen/qwen3.6-27b",
    "qwen/qwen3.8-27b",
]


def generate_inspection_narrative(
    report: dict,
    api_key: str | None = None,
    model: str = "openai/gpt-oss-120b",
    fallback: bool = True,
) -> tuple[str, str]:
    """Call Groq to produce a detailed inspection narrative from the VIN report.

    Returns (narrative_text, model_used).
    If ``fallback`` is True, tries FALLBACK_MODELS (preferred model first) on failure.
    """
    if Groq is None:
        raise RuntimeError("groq package is not installed. Run: pip install groq")

    key = api_key or os.environ.get("GROQ_API_KEY")
    if not key:
        raise ValueError(
            "Groq API key is required. Set GROQ_API_KEY environment variable "
            "or enter it in the Streamlit UI."
        )

    client = Groq(api_key=key)
    context = _build_context(report)
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Produce a detailed professional vehicle inspection report "
                "from this VIN-decoded and inspector data:\n\n" + context
            ),
        },
    ]

    candidates: list[str] = []
    for m in [model] + (FALLBACK_MODELS if fallback else []):
        if m and m not in candidates:
            candidates.append(m)

    errors: list[str] = []
    for candidate in candidates:
        try:
            response = client.chat.completions.create(
                model=candidate,
                messages=messages,
                temperature=0.25,
                max_tokens=4096,
            )
            content = response.choices[0].message.content
            if not content:
                errors.append(f"{candidate}: empty response")
                continue
            return content.strip(), candidate
        except Exception as e:
            errors.append(f"{candidate}: {e}")
            continue

    detail = " | ".join(errors) if errors else "unknown error"
    raise RuntimeError(f"All Groq models failed. Attempts: {detail}")


# ---------------------------------------------------------------------------
# Word document helpers
# ---------------------------------------------------------------------------

def _set_run_font(
    run,
    name: str = "Calibri",
    size_pt: int = 11,
    bold: bool = False,
    color: RGBColor | None = None,
    italic: bool = False,
):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _set_cell_shading(cell, hex_color: str) -> None:
    tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_margins(cell, top=40, bottom=40, left=60, right=60) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _add_hr(doc: Document) -> None:
    """Add a thin horizontal rule as a bottom-bordered empty paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _clean_md(text: str) -> str:
    """Strip residual markdown / escape artefacts from LLM output."""
    text = text.replace("\\*", "*").replace("\\_", "_")
    text = text.replace("\\-", "-")
    text = re.sub(r"^#+\s*", "", text)
    return text.strip()


def _add_rich_paragraph(doc: Document, text: str, size: int = 10, space_after: int = 4) -> None:
    """Add a paragraph, rendering **bold** segments."""
    text = _clean_md(text)
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = p.add_run(part[2:-2])
            _set_run_font(run, size_pt=size, bold=True, color=BLACK)
        else:
            run = p.add_run(part)
            _set_run_font(run, size_pt=size, color=BLACK)


def _add_bullet(doc: Document, text: str, size: int = 10) -> None:
    text = _clean_md(text)
    text = re.sub(r"^[\-\*\u2022]\s*", "", text)
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = p.add_run(part[2:-2])
            _set_run_font(run, size_pt=size, bold=True, color=BLACK)
        else:
            run = p.add_run(part)
            _set_run_font(run, size_pt=size, color=BLACK)


def _add_numbered(doc: Document, text: str, size: int = 10) -> None:
    text = _clean_md(text)
    text = re.sub(r"^\d+[.\)]\s*", "", text)
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = p.add_run(part[2:-2])
            _set_run_font(run, size_pt=size, bold=True, color=BLACK)
        else:
            run = p.add_run(part)
            _set_run_font(run, size_pt=size, color=BLACK)


def _kv_table(doc: Document, rows: list[tuple[str, str]], col_widths=(2.2, 4.6)) -> None:
    """Two-column key/value table with navy labels."""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (k, v) in enumerate(rows):
        cell0, cell1 = table.rows[i].cells
        cell0.width = Inches(col_widths[0])
        cell1.width = Inches(col_widths[1])
        _set_cell_margins(cell0)
        _set_cell_margins(cell1)
        if i % 2 == 1:
            _set_cell_shading(cell0, ROW_ALT_BG)
            _set_cell_shading(cell1, ROW_ALT_BG)
        p0 = cell0.paragraphs[0]
        run0 = p0.add_run(k)
        _set_run_font(run0, size_pt=9, bold=True, color=NAVY)
        p1 = cell1.paragraphs[0]
        run1 = p1.add_run(_v(v, "—"))
        _set_run_font(run1, size_pt=9, color=BLACK)
    doc.add_paragraph()


def _section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    _set_run_font(run, size_pt=12, bold=True, color=NAVY)
    _add_hr(doc)


def _parse_narrative(doc: Document, narrative_md: str) -> None:
    """Render LLM markdown narrative into clean Word paragraphs."""
    lines = narrative_md.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if re.match(r"^\|?\s*-+\s*\|", stripped) or stripped.startswith("|---"):
            i += 1
            continue

        if stripped.startswith("## "):
            _section_heading(doc, stripped[3:].strip())
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped[4:].strip())
            _set_run_font(run, size_pt=11, bold=True, color=STEEL)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_text = lines[i].strip()
                if re.match(r"^\|?\s*-+", row_text):
                    i += 1
                    continue
                cells = [c.strip() for c in row_text.strip("|").split("|")]
                if cells:
                    table_rows.append(cells)
                i += 1
            if table_rows and all(len(r) >= 2 for r in table_rows):
                start = 0
                first = " ".join(table_rows[0]).lower()
                if "item" in first or "specification" in first or "description" in first:
                    start = 1
                kv = []
                for r in table_rows[start:]:
                    k = _clean_md(r[0])
                    v = _clean_md(" — ".join(r[1:]))
                    if k:
                        kv.append((k, v))
                _kv_table(doc, kv)
            continue

        if re.match(r"^[\-\*\u2022]\s+", stripped) or (stripped.startswith("*") and not stripped.startswith("**")):
            if re.match(r"^\*?[A-Za-z][A-Za-z /&]+\*?$", stripped.replace("*", "").strip()) and len(stripped) < 40:
                label = stripped.replace("*", "").strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(label)
                _set_run_font(run, size_pt=10, bold=True, color=STEEL)
            else:
                _add_bullet(doc, stripped)
            i += 1
            continue

        if re.match(r"^\d+[.)]\s+", stripped):
            _add_numbered(doc, stripped)
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            label = stripped.strip("*").strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(label)
            _set_run_font(run, size_pt=10, bold=True, color=STEEL)
            i += 1
            continue

        _add_rich_paragraph(doc, stripped)
        i += 1


def _form_field_row(table, label: str, value: str = "", blank_hint: str = "") -> None:
    row = table.add_row()
    c0, c1 = row.cells
    _set_cell_margins(c0, top=30, bottom=30)
    _set_cell_margins(c1, top=30, bottom=30)
    p0 = c0.paragraphs[0]
    run = p0.add_run(label)
    _set_run_font(run, size_pt=9, bold=True, color=NAVY)
    p1 = c1.paragraphs[0]
    if value:
        run = p1.add_run(value)
        _set_run_font(run, size_pt=9, color=BLACK)
    else:
        hint = blank_hint or "________________________________"
        run = p1.add_run(hint)
        _set_run_font(run, size_pt=9, color=LIGHT_GRAY)


def _rating_row(doc: Document, label: str, current: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(f"{label}:  ")
    _set_run_font(run, size_pt=9, bold=True, color=NAVY)

    options = ["Excellent", "Good", "Fair", "Poor", "N/A"]
    cur = (current or "").lower().replace("_", " ")
    for opt in options:
        checked = "☑" if cur and opt.lower() in cur else "☐"
        run = p.add_run(f"  {checked} {opt}   ")
        _set_run_font(run, size_pt=9, color=BLACK if checked == "☑" else GRAY)


def _blank_lines(doc: Document, n: int = 3, label: str = "") -> None:
    if label:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run(label)
        _set_run_font(run, size_pt=9, bold=True, color=NAVY)
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("_" * 78)
        _set_run_font(run, size_pt=9, color=LIGHT_GRAY)


def _build_inspection_form(doc: Document, report: dict, branding: dict) -> None:
    """Append a formal multi-section vehicle inspection form (new page)."""
    doc.add_page_break()

    ident = report.get("identification") or {}
    engine = report.get("engine") or {}
    drive = report.get("drivetrain") or {}
    dims = report.get("dimensions") or {}
    mfg = report.get("manufacturing") or {}
    manual = report.get("manual") or {}
    cond = manual.get("condition") or {}
    wmi = report.get("wmi") or {}
    safety = report.get("safety") or {}

    vehicle_title = (
        " ".join(x for x in [ident.get("model_year"), ident.get("make"), ident.get("model")] if x)
        or "Vehicle"
    )
    company = branding.get("company") or branding.get("name") or "VIN Vehicle Report"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VEHICLE INSPECTION FORM")
    _set_run_font(run, size_pt=16, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company)
    _set_run_font(run, size_pt=10, color=GRAY)

    _add_hr(doc)

    _section_heading(doc, "A. Vehicle Identification (from VIN decode)")

    kv = [
        ("VIN", report.get("vin") or "—"),
        ("Year / Make / Model", vehicle_title),
        ("Trim / Series", f"{_v(ident.get('trim'), '—')} / {_v(ident.get('series'), '—')}"),
        ("Body Class / Type", f"{_v(ident.get('body_class'), '—')} / {_v(ident.get('vehicle_type'), '—')}"),
        ("Manufacturer", _v(ident.get("manufacturer"), "—")),
        ("Engine", f"{_v(engine.get('engine_model'), '—')} · {_v(engine.get('displacement'), '—')} · {_v(engine.get('cylinders'), '—')} cyl"),
        ("Fuel / Power", f"{_v(engine.get('fuel_type'), '—')} · {_v(engine.get('horsepower'), '—')} hp"),
        ("Transmission / Drive", f"{_v(drive.get('transmission'), '—')} · {_v(drive.get('drive_type'), '—')}"),
        ("Doors / GVWR", f"{_v(dims.get('doors'), '—')} · {_v(dims.get('gvwr'), '—')}"),
        ("Plant", f"{_v(mfg.get('plant_city'), '—')}, {_v(mfg.get('plant_state'), '—')}, {_v(mfg.get('plant_country'), '—')}"),
        ("WMI", f"{_v(wmi.get('code'), '—')} · {_v(wmi.get('country'), '—')}"),
        ("Report No.", report.get("id") or "—"),
        ("Data completeness", f"{report.get('data_completeness', 0)}%"),
    ]
    _kv_table(doc, kv)

    _section_heading(doc, "B. Client & Inspection Details")

    table = doc.add_table(rows=0, cols=2)
    table.autofit = False

    _form_field_row(table, "Client name", _v(manual.get("client_name"), ""))
    _form_field_row(table, "Client reference", _v(manual.get("client_reference"), ""))
    _form_field_row(table, "Registration / plate", _v(manual.get("registration"), ""))
    _form_field_row(table, "Odometer reading", _v(manual.get("mileage"), ""), blank_hint="_______________ km / mi")
    _form_field_row(table, "Purchase / asking price", _v(manual.get("purchase_price"), ""))
    _form_field_row(table, "Inspection date", _v(manual.get("inspection_date"), ""), blank_hint="____ / ____ / ________")
    _form_field_row(table, "Inspector name", _v(manual.get("inspector_name"), ""))
    _form_field_row(table, "Inspection location", "", blank_hint="________________________________")
    _form_field_row(table, "Weather / conditions", "", blank_hint="________________________________")
    doc.add_paragraph()

    _section_heading(doc, "C. Condition Ratings (tick one per row)")

    p = doc.add_paragraph()
    run = p.add_run(
        "Pre-filled from inspector notes where available. Tick or correct during physical inspection."
    )
    _set_run_font(run, size_pt=8, italic=True, color=GRAY)

    _rating_row(doc, "Overall condition", cond.get("overall", ""))
    _rating_row(doc, "Exterior body / paint", cond.get("exterior", ""))
    _rating_row(doc, "Interior / cabin", cond.get("interior", ""))
    _rating_row(doc, "Engine bay", cond.get("engine", ""))
    _rating_row(doc, "Transmission / drivetrain", cond.get("transmission", ""))
    _rating_row(doc, "Electrical / electronics", cond.get("electrical", ""))
    _rating_row(doc, "Tires / wheels", cond.get("tires", ""))
    _rating_row(doc, "Brakes", cond.get("brakes", ""))

    _section_heading(doc, "D. Physical Inspection Checklist")

    p = doc.add_paragraph()
    run = p.add_run("Mark each item:  ✓ Pass   ✗ Fail   — N/A   and note findings.")
    _set_run_font(run, size_pt=8, italic=True, color=GRAY)

    checklist_groups = [
        (
            "Exterior",
            [
                "Body panels alignment / gaps",
                "Paint condition / colour match / overspray",
                "Rust, corrosion, stone chips (wheel wells, sills, doors)",
                "Glass / windshield chips or cracks",
                "Lights, lenses, reflectors",
                "Mirrors, wipers, washers",
                "Doors, hood, trunk open/close / seals",
            ],
        ),
        (
            "Underbody / suspension",
            [
                "Frame / subframe / crossmembers",
                "Exhaust system condition & mounts",
                "Fuel / brake / power-steering lines",
                "Shock absorbers / struts / bushings",
                "CV joints / boots / axle seals",
                "Steering rack / linkage play",
            ],
        ),
        (
            "Engine bay",
            [
                "Oil / coolant / power-steering leaks",
                "Belts, hoses, clamps",
                "Battery terminals / corrosion / age",
                "Air filter / intake condition",
                "Fluid levels (oil, coolant, brake, washer)",
                "Evidence of aftermarket modifications",
            ],
        ),
        (
            "Interior",
            [
                "Seats, belts, airbags (visual)",
                "Dashboard, gauges, warning lamps",
                "HVAC operation / odours",
                "Power windows, locks, mirrors",
                "Infotainment / speakers",
                "Carpet / headliner water damage",
            ],
        ),
        (
            "Tires & brakes",
            [
                "Tread depth (all positions) mm/32nds",
                "Sidewall damage / uneven wear",
                "Brake pad thickness / disc condition",
                "Parking brake hold",
                "ABS / ESC warning status",
            ],
        ),
        (
            "Road test",
            [
                "Cold / hot start behaviour",
                "Idle quality / unusual noises",
                "Acceleration / turbo lag / hesitation",
                "Transmission shift quality",
                "Steering feel / pull / vibration",
                "Brake pedal feel / ABS activation",
                "HVAC performance under load",
            ],
        ),
    ]

    for group_name, items in checklist_groups:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(group_name)
        _set_run_font(run, size_pt=10, bold=True, color=STEEL)

        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(f"☐  {item}")
            _set_run_font(run, size_pt=9, color=BLACK)
            run = p.add_run("    Notes: _______________________________")
            _set_run_font(run, size_pt=8, color=LIGHT_GRAY)

    _section_heading(doc, "E. Safety Recalls (verify on site)")

    recalls = safety.get("recalls") or []
    if recalls:
        p = doc.add_paragraph()
        run = p.add_run(
            f"{len(recalls)} Year/Make/Model campaign(s) found in NHTSA data. "
            "Confirm exact VIN applicability at https://www.nhtsa.gov/recalls"
        )
        _set_run_font(run, size_pt=9, color=BLACK)
        for r in recalls[:6]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(
                f"{r.get('campaign_number', '')} — {r.get('component') or 'Component N/A'}"
            )
            _set_run_font(run, size_pt=9, color=BLACK)
        p = doc.add_paragraph()
        run = p.add_run("Recall status checked on site:  ☐ Open  ☐ Closed / remedied  ☐ N/A")
        _set_run_font(run, size_pt=9, color=BLACK)
    else:
        p = doc.add_paragraph()
        run = p.add_run(
            "No Year/Make/Model recalls returned from free NHTSA source, or data unavailable. "
            "Always re-check with the exact VIN before delivery."
        )
        _set_run_font(run, size_pt=9, color=BLACK)

    _section_heading(doc, "F. Inspector Notes & Findings")

    existing_notes = (manual.get("notes") or "").strip()
    existing_findings = (manual.get("additional_findings") or "").strip()
    if existing_notes or existing_findings:
        p = doc.add_paragraph()
        run = p.add_run("Previously recorded notes:")
        _set_run_font(run, size_pt=9, bold=True, color=NAVY)
        if existing_notes:
            _add_rich_paragraph(doc, existing_notes, size=9)
        if existing_findings:
            _add_rich_paragraph(doc, existing_findings, size=9)

    _blank_lines(doc, 4, "Additional findings from physical inspection:")

    _section_heading(doc, "G. Summary Recommendation")

    p = doc.add_paragraph()
    run = p.add_run("Overall recommendation (tick one):")
    _set_run_font(run, size_pt=9, bold=True, color=NAVY)

    for opt in [
        "☐  Suitable for purchase / use as-is",
        "☐  Suitable subject to minor repairs (list below)",
        "☐  Major repairs required before use / sale",
        "☐  Not recommended without further specialist assessment",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(opt)
        _set_run_font(run, size_pt=9, color=BLACK)

    _blank_lines(doc, 3, "Required repairs / follow-up:")

    _section_heading(doc, "H. Signatures")

    sig = doc.add_table(rows=3, cols=2)
    sig.autofit = False
    labels = [
        ("Inspector name", "Client / owner name"),
        ("Signature", "Signature"),
        ("Date", "Date"),
    ]
    for i, (left, right) in enumerate(labels):
        c0, c1 = sig.rows[i].cells
        _set_cell_margins(c0, top=60, bottom=60)
        _set_cell_margins(c1, top=60, bottom=60)
        for cell, lab in ((c0, left), (c1, right)):
            p = cell.paragraphs[0]
            run = p.add_run(f"{lab}:\n\n______________________________")
            _set_run_font(run, size_pt=9, color=GRAY)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "This form documents a visual / operational inspection only. It is not a guarantee of "
        "mechanical condition, title status, or accident history. Always verify open recalls "
        "at https://www.nhtsa.gov/recalls using the exact VIN."
    )
    _set_run_font(run, size_pt=8, italic=True, color=GRAY)


def build_inspection_docx(
    report: dict,
    narrative_md: str,
    branding: dict | None = None,
) -> bytes:
    """Professional multi-page Word report:

    Page 1+: AI narrative (cleanly formatted)
    Following page(s): formal Vehicle Inspection Form (prefilled + blanks)
    """
    branding = branding or {}
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    styles = doc.styles
    for style_name in ("Normal", "List Bullet", "List Number"):
        try:
            styles[style_name].font.name = "Calibri"
            styles[style_name].font.size = Pt(10)
        except KeyError:
            pass

    ident = report.get("identification") or {}
    vehicle_title = (
        " ".join(x for x in [ident.get("model_year"), ident.get("make"), ident.get("model")] if x)
        or "Vehicle Inspection Report"
    )
    company = branding.get("company") or branding.get("name") or "VIN Vehicle Report"
    title_text = (branding.get("default_report_title") or "Detailed Vehicle Inspection Report").upper()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title_text)
    _set_run_font(run, size_pt=18, bold=True, color=NAVY)

    p = doc.add_paragraph()
    run = p.add_run(company)
    _set_run_font(run, size_pt=11, bold=True, color=BLACK)

    contact_bits = [x for x in [branding.get("email"), branding.get("phone"), branding.get("website")] if x]
    if contact_bits:
        p = doc.add_paragraph()
        run = p.add_run("  ·  ".join(contact_bits))
        _set_run_font(run, size_pt=9, color=GRAY)

    _add_hr(doc)

    p = doc.add_paragraph()
    run = p.add_run(vehicle_title)
    _set_run_font(run, size_pt=14, bold=True, color=STEEL)

    p = doc.add_paragraph()
    run = p.add_run(f"VIN: {report.get('vin', '')}")
    _set_run_font(run, size_pt=11, bold=True, color=BLACK)

    meta = (
        f"Report No: {report.get('id', '')}    ·    "
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}    ·    "
        f"Data completeness: {report.get('data_completeness', 0)}%"
    )
    p = doc.add_paragraph()
    run = p.add_run(meta)
    _set_run_font(run, size_pt=9, color=GRAY)

    p = doc.add_paragraph()
    run = p.add_run(
        "Part 1 — AI-assisted narrative compiled from free public NHTSA data and inspector notes. "
        "Part 2 — Formal inspection form for on-site completion."
    )
    _set_run_font(run, size_pt=8, italic=True, color=GRAY)

    _add_hr(doc)

    _parse_narrative(doc, narrative_md)

    doc.add_paragraph()
    _add_hr(doc)
    p = doc.add_paragraph()
    run = p.add_run("Important Limitations")
    _set_run_font(run, size_pt=10, bold=True, color=NAVY)

    limits = [
        "Generated from free public NHTSA data and optional inspector notes. Does not include ownership history, accident records, title brands, odometer history, service records, or auction data.",
        "Absence of information does not confirm a clean history.",
        "Always verify open recalls at https://www.nhtsa.gov/recalls using the exact VIN.",
        "This document is not a substitute for a professional mechanical inspection or a paid vehicle-history report.",
    ]
    for lim in limits:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(lim)
        _set_run_font(run, size_pt=8, color=GRAY)

    disc = branding.get("disclaimer") or (
        "Information in this report is compiled from available external data sources and user-provided information. "
        "Availability and accuracy depend on the underlying sources. This report is not a substitute for a professional "
        "mechanical inspection, official title verification, or a comprehensive paid vehicle-history report."
    )
    p = doc.add_paragraph()
    run = p.add_run("Disclaimer: ")
    _set_run_font(run, size_pt=8, bold=True, color=GRAY)
    run = p.add_run(disc)
    _set_run_font(run, size_pt=8, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(branding.get("report_footer") or "Confidential – For client use only")
    _set_run_font(run, size_pt=8, color=LIGHT_GRAY)

    _build_inspection_form(doc, report, branding)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
