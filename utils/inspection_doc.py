"""Generate a detailed vehicle inspection Word report via Groq LLM + python-docx."""

from __future__ import annotations

import json
import os
from datetime import datetime
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

NA = "Not available from free sources"


def _v(val: Any, fallback: str = NA) -> str:
    if val is None:
        return fallback
    s = str(val).strip()
    return s if s else fallback


def _set_run_font(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(30, 58, 95)
    return h


def _kv_table(doc: Document, pairs: list):
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(pairs):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
    doc.add_paragraph()


def build_report_context(report: dict, branding: dict) -> str:
    ident = report.get("identification") or {}
    engine = report.get("engine") or {}
    drive = report.get("drivetrain") or {}
    dims = report.get("dimensions") or {}
    mfg = report.get("manufacturing") or {}
    safety = report.get("safety") or {}
    manual = report.get("manual") or {}
    wmi = report.get("wmi") or {}
    complaints = report.get("complaints") or {}
    ncap = report.get("ncap") or {}
    epa = report.get("epa") or {}

    payload = {
        "vin": report.get("vin"),
        "report_id": report.get("id"),
        "identification": ident,
        "engine": engine,
        "drivetrain": drive,
        "dimensions": dims,
        "manufacturing": mfg,
        "wmi": wmi,
        "recalls_count": len(safety.get("recalls") or []),
        "recalls_note": safety.get("recalls_note"),
        "sample_recalls": (safety.get("recalls") or [])[:5],
        "complaints_count": complaints.get("count"),
        "complaints_crash_flags_in_sample": complaints.get("crash_count"),
        "sample_complaint_components": [
            c.get("components") for c in (complaints.get("items") or [])[:8]
        ],
        "ncap": {
            "overall": ncap.get("overall"),
            "frontal": ncap.get("frontal"),
            "side": ncap.get("side"),
            "rollover": ncap.get("rollover"),
            "description": ncap.get("vehicle_description"),
        },
        "epa": {
            "city": epa.get("city"),
            "highway": epa.get("highway"),
            "combined": epa.get("combined"),
            "fuel_type": epa.get("fuel_type"),
            "annual_fuel_cost": epa.get("annual_fuel_cost"),
            "option": epa.get("option_text"),
        },
        "manual_inspection": manual,
        "data_completeness_pct": report.get("data_completeness"),
        "inspector_branding": {
            "name": branding.get("name"),
            "company": branding.get("company"),
        },
    }
    return json.dumps(payload, indent=2, ensure_ascii=False)


SYSTEM_PROMPT = """You are an experienced automotive inspector writing a professional vehicle inspection narrative for a client report.

STRICT RULES:
1. Use ONLY facts present in the provided JSON context. Never invent accident history, ownership, title status, odometer fraud, or service records.
2. If a field is missing, say it was not available from free public sources — do not guess.
3. Recalls and complaints are Year/Make/Model level, not proof of this exact VIN's history.
4. Write in clear professional English suitable for a client-facing inspection report.
5. Structure your response with these exact section headings (markdown ##):

## Executive Summary
## Vehicle Identification
## Factory Specifications
## Powertrain and Drivetrain
## Manufacturing Origin
## Safety Ratings (NCAP)
## Fuel Economy (EPA)
## Safety Recalls Discussion
## Owner Complaints Patterns
## Physical Inspection Findings
## Condition Assessment
## Limitations and Disclaimers
## Recommendations

6. In Physical Inspection Findings and Condition Assessment, base content on the manual_inspection fields. If not assessed, state that clearly.
7. Recommendations should be practical without claiming clean history.
8. Be detailed but factual.
"""


def generate_narrative_with_groq(report: dict, branding: dict, api_key: str, model: str = "llama-3.3-70b-versatile") -> str:
    try:
        from groq import Groq
    except ImportError as e:
        raise RuntimeError("groq package not installed. Run: pip install groq") from e

    client = Groq(api_key=api_key)
    context = build_report_context(report, branding)
    user_msg = (
        "Write a detailed professional vehicle inspection report narrative from this data.\n\n"
        f"CONTEXT JSON:\n{context}"
    )

    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.3,
        max_tokens=6000,
    )
    text = completion.choices[0].message.content or ""
    if not text.strip():
        raise RuntimeError("Groq returned an empty response.")
    return text.strip()


def _add_md_like_body(doc: Document, narrative: str):
    for block in narrative.split("\n"):
        line = block.rstrip()
        if not line:
            continue
        if line.startswith("## "):
            _add_heading(doc, line[3:].strip(), level=1)
        elif line.startswith("### "):
            _add_heading(doc, line[4:].strip(), level=2)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = "Calibri"


def build_inspection_docx(report: dict, branding: dict, narrative: str) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    ident = report.get("identification") or {}
    vehicle_title = " ".join(
        x for x in [ident.get("model_year"), ident.get("make"), ident.get("model")] if x
    ) or "Vehicle Inspection Report"

    title = doc.add_heading(branding.get("default_report_title") or "Vehicle Inspection Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 58, 95)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(vehicle_title)
    _set_run_font(r, 14, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        f"VIN: {report.get('vin', '')}  |  Report: {report.get('id', '')}  |  "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    _set_run_font(r, 10, color=(80, 80, 80))

    company = branding.get("company") or branding.get("name") or ""
    if company:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(company)
        _set_run_font(r, 11, bold=True)

    doc.add_paragraph()
    _add_heading(doc, "Quick Facts (Decoded Data)", level=1)
    engine = report.get("engine") or {}
    drive = report.get("drivetrain") or {}
    mfg = report.get("manufacturing") or {}
    _kv_table(
        doc,
        [
            ("VIN", _v(report.get("vin"))),
            ("Make / Model / Year", vehicle_title),
            ("Trim", _v(ident.get("trim"))),
            ("Body Class", _v(ident.get("body_class"))),
            ("Engine", _v(" \u00b7 ".join(x for x in [
                f"{engine.get('cylinders')} cyl" if engine.get("cylinders") else None,
                engine.get("displacement"),
                engine.get("fuel_type"),
            ] if x))),
            ("Drive", _v(drive.get("drive_type"))),
            ("Plant", _v(", ".join(x for x in [mfg.get("plant_city"), mfg.get("plant_country")] if x))),
            ("Data completeness", f"{report.get('data_completeness', 0)}%"),
        ],
    )

    _add_md_like_body(doc, narrative)

    _add_heading(doc, "Document Disclaimer", level=1)
    disc = branding.get("disclaimer") or (
        "This document combines free public data (primarily NHTSA and EPA) with inspector-provided notes. "
        "It is not a substitute for a mechanical inspection, title search, or paid vehicle-history report. "
        "Absence of information does not mean an event did not occur."
    )
    p = doc.add_paragraph(disc)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(60, 60, 60)

    footer = branding.get("report_footer") or "Confidential \u2013 For client use only"
    p = doc.add_paragraph(footer)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_inspection_docx(
    report: dict,
    branding: dict,
    api_key: str | None = None,
    model: str = "llama-3.3-70b-versatile",
) -> bytes:
    key = api_key or os.environ.get("GROQ_API_KEY")
    if not key:
        raise RuntimeError(
            "Groq API key missing. Set GROQ_API_KEY environment variable or enter it in the app."
        )
    narrative = generate_narrative_with_groq(report, branding, key, model=model)
    return build_inspection_docx(report, branding, narrative)
